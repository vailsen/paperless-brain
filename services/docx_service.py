# services/docx_service.py
"""Generate DIN 5008 letters from a DOCX template.

This is deliberately a German-market feature: DIN 5008 is the German business-
letter norm, so the template, the DD.MM.YYYY date, the default closing
("Mit freundlichen Grüßen") and labels like "Bezug:" stay German regardless of
the UI language. The letter *content* (subject, salutation, body) comes from
the LLM in whatever language the user requested.
"""

import io
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

_TEMPLATE = Path(__file__).parent.parent / "config" / "templates" / "din5008_template.docx"
_BODY_SLOT_RE = re.compile(r"\{\{BODY_PARAGRAPH\}\}")


def _replace_para(para, mapping: dict) -> None:
    """Replace {{KEY}} placeholders in a paragraph, handling split runs."""
    full_text = "".join(r.text for r in para.runs)
    changed = False
    for key, value in mapping.items():
        token = f"{{{{{key}}}}}"
        if token in full_text:
            full_text = full_text.replace(token, str(value))
            changed = True
    if changed and para.runs:
        para.runs[0].text = full_text
        for r in para.runs[1:]:
            r.text = ""


def _all_paragraphs(doc: Document):
    """Yield all paragraphs: body + every section's header and footer."""
    yield from doc.paragraphs
    for section in doc.sections:
        if not section.header.is_linked_to_previous:
            yield from section.header.paragraphs
        if not section.footer.is_linked_to_previous:
            yield from section.footer.paragraphs


def _replace_body_paragraphs(doc: Document, body_paras: list[str]) -> None:
    """Replace single {{BODY_PARAGRAPH}} slot with first para, insert rest after it."""
    slot = next(
        (p for p in doc.paragraphs
         if _BODY_SLOT_RE.search("".join(r.text for r in p.runs))),
        None,
    )
    if slot is None:
        return

    ref_style = slot.style
    anchor = slot._element

    _replace_para(slot, {"BODY_PARAGRAPH": body_paras[0] if body_paras else ""})
    slot.paragraph_format.space_after = Pt(6)

    last_para = slot
    insert_after = anchor
    for text in body_paras[1:]:
        new_para = doc.add_paragraph(text, style=ref_style)
        new_para.paragraph_format.space_after = Pt(6)
        p_elem = new_para._element
        p_elem.getparent().remove(p_elem)
        insert_after.addnext(p_elem)
        insert_after = p_elem
        last_para = new_para

    last_para.paragraph_format.space_after = Pt(12)


def _collapse_extra_sections(doc: Document) -> None:
    """Remove embedded sectPr breaks that produce blank continuation pages."""
    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                pPr.remove(sectPr)


def generate_letter_docx(
    sender: dict,
    recipient: dict,
    subject: str,
    salutation: str,
    body_paras: list[str],
    closing: str = "Mit freundlichen Grüßen",
    src_doc_info: str = "",
    source_cross_ref: str = "",
    reference_doc_id: int | None = None,
) -> bytes:
    """
    Fill the DIN 5008 template and return the DOCX as bytes.

    sender keys: name, street, plz, city, phone, email, company
    recipient keys: name, street, postcode, city
    source_cross_ref: "Ihr Zeichen" — e.g. "Rechn. Nr. 260412"
    src_doc_info: "Bezug" line — short title of referenced doc (optional)
    """
    doc = Document(_TEMPLATE)

    sender_name = sender.get("name", "")
    if sender.get("company"):
        sender_name = f"{sender_name}, {sender['company']}"

    cross_ref = source_cross_ref or ""
    if not cross_ref and reference_doc_id:
        cross_ref = f"Dok. #{reference_doc_id}"

    base_mapping = {
        "SENDER_NAME":          sender_name,
        "SENDER_STREET":        sender.get("street", ""),
        "SENDER_PLZ":           sender.get("plz", ""),
        "SENDER_CITY":          sender.get("city", ""),
        "SENDER_PHONE":         sender.get("phone", ""),
        "SENDER_EMAIL":         sender.get("email", ""),
        "RECIPIENT_NAME":       recipient.get("name", ""),
        "RECIPIENT_STREET":     recipient.get("street", ""),
        "RECIPIENT_PLZ":        recipient.get("postcode", ""),
        "RECIPIENT_CITY":       recipient.get("city", ""),
        "DATE":                 date.today().strftime("%d.%m.%Y"),
        "SOURCE_CROSS_REF":     cross_ref,
        "SUBJECT":              subject,
        "SALUTATION":           salutation,
        "CLOSING":              closing,
        "SOURCE_DOCUMENT_INFO": src_doc_info,
    }

    # Replace all non-body placeholders (body + headers + footers)
    for para in _all_paragraphs(doc):
        _replace_para(para, base_mapping)

    # Flexible body paragraphs — any count, any number of template slots
    _replace_body_paragraphs(doc, body_paras)

    # Remove "Bezug:" paragraph when src_doc_info is empty
    if not src_doc_info:
        for para in doc.paragraphs:
            if _BODY_SLOT_RE.search(para.text) or (
                "Bezug:" in para.text
                and not para.text.replace("Bezug:", "").strip()
            ):
                para._element.getparent().remove(para._element)
                break

    _collapse_extra_sections(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
